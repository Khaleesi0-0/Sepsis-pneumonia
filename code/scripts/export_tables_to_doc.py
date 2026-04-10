from __future__ import annotations

import importlib
import subprocess
import sys
from pathlib import Path

import pandas as pd


ROOT = Path(__file__).resolve().parents[2]
TABLE_ROOT = ROOT / "results" / "tables"
DOC_PATH = TABLE_ROOT / "tables_export.docx"

TABLE_SOURCES = [
    TABLE_ROOT / "overall_summary_table.csv",
    TABLE_ROOT / "outcome_year_summary_table.csv",
    TABLE_ROOT / "period_summary" / "all_outcomes_period_summary_table.csv",
    TABLE_ROOT / "period_summary" / "sepsis_period_summary_table.csv",
    TABLE_ROOT / "period_summary" / "pneumonia_period_summary_table.csv",
    TABLE_ROOT / "period_summary" / "combined_period_summary_table.csv",
]


def _ensure_python_docx():
    try:
        return importlib.import_module("docx")
    except ModuleNotFoundError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        return importlib.import_module("docx")


docx = _ensure_python_docx()
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, *, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = 0
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)


def _style_table(table) -> None:
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9)
        if row_idx == 0:
            for cell in row.cells:
                _set_cell_shading(cell, "D9E2F3")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def _set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)


def _add_dataframe_table(document: Document, title: str, df: pd.DataFrame) -> None:
    document.add_heading(title, level=1)
    table = document.add_table(rows=1, cols=len(df.columns))
    table.autofit = False

    for col_idx, column_name in enumerate(df.columns):
        _set_cell_text(table.rows[0].cells[col_idx], column_name, bold=True, size=9)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for col_idx, value in enumerate(row.tolist()):
            _set_cell_text(cells[col_idx], "" if pd.isna(value) else value, size=8)

    _style_table(table)
    document.add_paragraph("")


def _load_table(path: Path) -> pd.DataFrame:
    return pd.read_csv(path)


def export_tables() -> None:
    document = Document()
    _set_landscape(document.sections[0])

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(9)

    title = document.add_paragraph()
    title.alignment = 1
    run = title.add_run("Summary Tables")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    subtitle = document.add_paragraph()
    subtitle.alignment = 1
    run = subtitle.add_run("Publication-ready export of all generated tables")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    for table_path in TABLE_SOURCES:
        if not table_path.exists():
            continue
        df = _load_table(table_path)
        _add_dataframe_table(document, table_path.stem.replace("_", " ").title(), df)
        document.add_page_break()

    DOC_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(DOC_PATH)


if __name__ == "__main__":
    export_tables()
