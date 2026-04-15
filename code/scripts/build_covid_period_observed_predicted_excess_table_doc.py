from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
TABLE_DIR = ROOT / "results" / "tables"

AAMR_SOURCE = TABLE_DIR / "monthly_aamr_predicted_2020_2025.csv"
DEATHS_SOURCE = TABLE_DIR / "monthly_deaths_predicted_2020_2025.csv"

CSV_OUTPUT = TABLE_DIR / "covid_period_observed_predicted_excess_table.csv"
DOC_OUTPUT = TABLE_DIR / "covid_period_observed_predicted_excess_table.docx"
YEARLY_OUTPUT = TABLE_DIR / "covid_yearly_observed_predicted_excess_values_2020_2023.csv"

YEAR_COL = "Year"
DISEASE_COL = "Disease"

DISEASE_ORDER = ["Pneumonia", "Pneumonia/ARDS", "Pneumonia/Sepsis"]
DISEASE_NAME_MAP = {
    "ARDS (ARDS + Pneumonia)": "Pneumonia/ARDS",
    "ARDS/Pneumonia": "Pneumonia/ARDS",
    "Sepsis+ Pneumonia": "Pneumonia/Sepsis",
    "Sepsis/Pneumonia": "Pneumonia/Sepsis",
    "Combined": "Pneumonia/Sepsis",
    "combined": "Pneumonia/Sepsis",
}

COVID_YEARS = [2020, 2021, 2022, 2023]


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, *, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = 0
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)


def _set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)


def _prepare(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out[DISEASE_COL] = out[DISEASE_COL].astype("string").str.strip().replace(DISEASE_NAME_MAP)
    out[YEAR_COL] = pd.to_numeric(out[YEAR_COL], errors="coerce")
    out = out[out[DISEASE_COL].isin(DISEASE_ORDER)].copy()
    return out


def _build_yearly_values_df(aamr_df: pd.DataFrame, deaths_df: pd.DataFrame) -> pd.DataFrame:
    aamr = aamr_df[aamr_df[YEAR_COL].isin(COVID_YEARS)].copy()
    deaths = deaths_df[deaths_df[YEAR_COL].isin(COVID_YEARS)].copy()

    for col in ["Age Adjusted Rate", "Predicted AAMR", "Actual - Predicted"]:
        aamr[col] = pd.to_numeric(aamr[col], errors="coerce")
    for col in ["Deaths", "Predicted Deaths", "Actual - Predicted Deaths"]:
        deaths[col] = pd.to_numeric(deaths[col], errors="coerce")

    annual_aamr = (
        aamr.groupby([DISEASE_COL, YEAR_COL], as_index=False)[
            ["Age Adjusted Rate", "Predicted AAMR", "Actual - Predicted"]
        ]
        .sum(min_count=1)
    )
    annual_deaths = (
        deaths.groupby([DISEASE_COL, YEAR_COL], as_index=False)[
            ["Deaths", "Predicted Deaths", "Actual - Predicted Deaths"]
        ]
        .sum(min_count=1)
    )

    annual = annual_aamr.merge(annual_deaths, on=[DISEASE_COL, YEAR_COL], how="outer")
    annual = annual.rename(
        columns={
            "Age Adjusted Rate": "Observed AAMR",
            "Predicted AAMR": "Predicted AAMR",
            "Actual - Predicted": "Excess AAMR",
            "Deaths": "Observed Death",
            "Predicted Deaths": "Predicted Death",
            "Actual - Predicted Deaths": "Excess Death",
        }
    )
    annual[DISEASE_COL] = pd.Categorical(annual[DISEASE_COL], categories=DISEASE_ORDER, ordered=True)
    annual = annual.sort_values([DISEASE_COL, YEAR_COL]).reset_index(drop=True)
    return annual


def _build_table_df(yearly_df: pd.DataFrame) -> pd.DataFrame:
    rows: list[dict[str, object]] = []
    numeric_cols = [
        "Observed AAMR",
        "Predicted AAMR",
        "Excess AAMR",
        "Observed Death",
        "Predicted Death",
        "Excess Death",
    ]

    for disease in DISEASE_ORDER:
        disease_df = yearly_df[yearly_df[DISEASE_COL].eq(disease)].copy()
        for year in COVID_YEARS:
            year_row = disease_df[disease_df[YEAR_COL].eq(year)]
            row = {"Disease": disease, "Year": str(year)}
            for col in numeric_cols:
                value = year_row[col].iloc[0] if not year_row.empty else pd.NA
                row[col] = value
            rows.append(row)

        total_row = {"Disease": disease, "Year": "Total"}
        for col in numeric_cols:
            series = pd.to_numeric(disease_df[col], errors="coerce")
            if "AAMR" in col:
                total_row[col] = series.mean()
            else:
                total_row[col] = series.sum(min_count=1)
        rows.append(total_row)

    out = pd.DataFrame(rows)
    for col in numeric_cols:
        out[col] = pd.to_numeric(out[col], errors="coerce").map(
            lambda x: "" if pd.isna(x) else f"{x:.2f}"
        )
    return out[["Disease", "Year"] + numeric_cols]


def _add_table(document: Document, df: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.autofit = True

    for i, col in enumerate(df.columns):
        _set_cell_text(table.rows[0].cells[i], col, bold=True, size=9)
        _set_cell_shading(table.rows[0].cells[i], "D9E2F3")

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            _set_cell_text(cells[i], row[col], size=8)


def export_doc(df: pd.DataFrame) -> None:
    doc = Document()
    _set_landscape(doc.sections[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = 1
    run = title.add_run("COVID-Year Observed, Predicted, and Excess Mortality by Disease")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    subtitle = doc.add_paragraph()
    subtitle.alignment = 1
    run = subtitle.add_run(
        "Rows: Pneumonia, Pneumonia/ARDS, Pneumonia/Sepsis by year (2020-2023) plus Total"
    )
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    _add_table(doc, df)

    note = doc.add_paragraph()
    note_run = note.add_run(
        "Abbreviation: AAMR, age-adjusted mortality rate per 100,000. "
        "Yearly values are aggregated from monthly values; Total uses sum for deaths and mean for AAMR."
    )
    note_run.italic = True
    note_run.font.name = "Times New Roman"
    note_run.font.size = Pt(9)

    DOC_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOC_OUTPUT)


def build_and_export() -> None:
    if not AAMR_SOURCE.exists():
        raise FileNotFoundError(f"Missing AAMR source table: {AAMR_SOURCE}")
    if not DEATHS_SOURCE.exists():
        raise FileNotFoundError(f"Missing deaths source table: {DEATHS_SOURCE}")

    aamr_df = _prepare(pd.read_csv(AAMR_SOURCE))
    deaths_df = _prepare(pd.read_csv(DEATHS_SOURCE))

    yearly_df = _build_yearly_values_df(aamr_df, deaths_df)
    for col in [
        "Observed AAMR",
        "Predicted AAMR",
        "Excess AAMR",
        "Observed Death",
        "Predicted Death",
        "Excess Death",
    ]:
        yearly_df[col] = pd.to_numeric(yearly_df[col], errors="coerce").round(2)
    yearly_df.to_csv(YEARLY_OUTPUT, index=False)

    out_df = _build_table_df(yearly_df)
    out_df.to_csv(CSV_OUTPUT, index=False)
    export_doc(out_df)


if __name__ == "__main__":
    build_and_export()
