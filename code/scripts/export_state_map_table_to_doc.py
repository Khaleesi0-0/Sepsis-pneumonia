from __future__ import annotations

import importlib
import subprocess
import sys
from pathlib import Path

import pandas as pd


ROOT = Path(__file__).resolve().parents[2]
TABLE_DIR = ROOT / "results" / "tables"
CSV_PATH = TABLE_DIR / "state_map_period_data.csv"
DOC_PATH = TABLE_DIR / "state_map_period_data.docx"

DISEASE_ORDER = ["Pneumonia", "Pneumonia/ARDS", "Pneumonia/Sepsis"]
DISEASE_NAME_MAP = {
    "ARDS (ARDS + Pneumonia)": "Pneumonia/ARDS",
    "ARDS/Pneumonia": "Pneumonia/ARDS",
    "Sepsis+ Pneumonia": "Pneumonia/Sepsis",
    "Sepsis/Pneumonia": "Pneumonia/Sepsis",
    "Combined": "Pneumonia/Sepsis",
    "combined": "Pneumonia/Sepsis",
}


def _ensure_python_docx():
    try:
        return importlib.import_module("docx")
    except ModuleNotFoundError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        return importlib.import_module("docx")


docx = _ensure_python_docx()
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def _set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, *, bold: bool = False, size: float = 7.0) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)


def _repeat_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _format_table_values(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    outcome_col = "Outcome" if "Outcome" in out.columns else ("outcome" if "outcome" in out.columns else None)
    if outcome_col is not None:
        out[outcome_col] = out[outcome_col].astype("string").str.strip().replace(DISEASE_NAME_MAP)
        out[outcome_col] = pd.Categorical(out[outcome_col], categories=DISEASE_ORDER, ordered=True)

    if "Period" in out.columns:
        period_order = {
            "Pre-pandemic (2010–2019)": 1,
            "Pandemic (2020-2023)": 2,
            "Pandemic (2020–2023)": 2,
            "Post-pandemic (2024–2025)": 3,
        }
        out["__period_ord"] = out["Period"].astype("string").map(period_order)
    else:
        out["__period_ord"] = pd.NA

    for col in ["Deaths", "Population"]:
        if col in out.columns:
            out[col] = pd.to_numeric(out[col], errors="coerce").round(0).astype("Int64").astype("string")
    for col in ["Crude rate (/100,000)", "AAMR (/100,000)"]:
        if col in out.columns:
            out[col] = pd.to_numeric(out[col], errors="coerce").map(lambda value: f"{value:.2f}" if pd.notna(value) else "")
    if "State Code" in out.columns:
        out["State Code"] = pd.to_numeric(out["State Code"], errors="coerce").round(0).astype("Int64").astype("string")
    sort_cols = []
    if outcome_col is not None:
        sort_cols.append(outcome_col)
    sort_cols.extend(["__period_ord"])
    if "State" in out.columns:
        sort_cols.append("State")
    out = out.sort_values(sort_cols, na_position="last").reset_index(drop=True)
    out = out.drop(columns=["__period_ord"], errors="ignore")
    return out.fillna("")


def export_state_map_table() -> None:
    if not CSV_PATH.exists():
        raise FileNotFoundError(f"Missing table: {CSV_PATH}")

    df = _format_table_values(pd.read_csv(CSV_PATH))

    document = Document()
    _set_landscape(document.sections[0])

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(8)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("State-Level Mortality Data Used for Map Figure")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Period-aggregated deaths, population, crude mortality rates, and AAMR by outcome and state")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)

    table = document.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    header = table.rows[0]
    _repeat_header_row(header)
    for col_idx, column_name in enumerate(df.columns):
        cell = header.cells[col_idx]
        _set_cell_text(cell, column_name, bold=True, size=7.3)
        _set_cell_shading(cell, "D9E2F3")

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for col_idx, value in enumerate(row.tolist()):
            _set_cell_text(cells[col_idx], value, size=6.7)

    DOC_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(DOC_PATH)


if __name__ == "__main__":
    export_state_map_table()
