from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
CLEANED_DIR = ROOT / "data" / "cleaned"
SEX_DIR = CLEANED_DIR / "disease_sex"
TABLE_DIR = ROOT / "results" / "tables"
DOC_OUTPUT_PATH = TABLE_DIR / "yearly_age_sex_aamr_tables.docx"

YEAR_COL = "Year"
NOTES_COL = "Notes"
SEX_COL = "Sex"
AGE_GROUP_COL = "Age Group"
DEATHS_COL = "Deaths"
POP_COL = "Population"
SEX_AAMR_COL = "Age Adjusted Rate"
AGE_AAMR_COL = "AAMR"

SEX_ORDER = ["Female", "Male"]
AGE_ORDER = ["<25", "25-44", "45-64", "65+"]

DATASETS = {
    "Pneumonia": {
        "sex": "pneumonia_sex.csv",
        "age": CLEANED_DIR / "pneumonia_age.csv",
    },
    "Pneumonia/ARDS": {
        "sex": "ards_sex.csv",
        "age": CLEANED_DIR / "ards_age.csv",
    },
    "Pneumonia/Sepsis": {
        "sex": "combined_sex.csv",
        "age": CLEANED_DIR / "combined_age.csv",
    },
}


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _resolve_sex_path(file_name: str) -> Path:
    primary = SEX_DIR / file_name
    fallback = CLEANED_DIR / file_name
    if primary.exists():
        return primary
    if fallback.exists():
        return fallback
    raise FileNotFoundError(
        f"Sex dataset not found: {file_name}. Checked {primary} and {fallback}."
    )


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: int = 9,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)


def _configure_section(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)


def _normalize_year_column(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out[YEAR_COL] = (
        out[YEAR_COL]
        .astype("string")
        .str.strip()
        .str.replace(r"\s*\(provisional\)$", "", regex=True)
    )
    out[YEAR_COL] = pd.to_numeric(out[YEAR_COL], errors="coerce")
    return out


def _aggregate_deaths_and_weighted_aamr(
    df: pd.DataFrame,
    group_cols: list[str],
    *,
    deaths_col: str,
    rate_col: str,
    pop_col: str,
) -> pd.DataFrame:
    out = df.copy()
    valid_rate = out[rate_col].notna() & out[pop_col].notna() & (out[pop_col] > 0)
    out["__aamr_num"] = 0.0
    out.loc[valid_rate, "__aamr_num"] = out.loc[valid_rate, rate_col] * out.loc[valid_rate, pop_col]
    out["__aamr_den"] = 0.0
    out.loc[valid_rate, "__aamr_den"] = out.loc[valid_rate, pop_col]

    grouped = (
        out.groupby(group_cols, as_index=False, dropna=False)
        .agg(
            Deaths=(deaths_col, "sum"),
            Population=(pop_col, "sum"),
            __aamr_num=("__aamr_num", "sum"),
            __aamr_den=("__aamr_den", "sum"),
        )
    )
    grouped["AAMR"] = grouped["__aamr_num"] / grouped["__aamr_den"]
    grouped.loc[grouped["__aamr_den"] <= 0, "AAMR"] = pd.NA
    return grouped.drop(columns=["__aamr_num", "__aamr_den"])


def _weighted_rate_from_grouped(df: pd.DataFrame) -> float | None:
    if df.empty:
        return None
    pop = pd.to_numeric(df["Population"], errors="coerce")
    rate = pd.to_numeric(df["AAMR"], errors="coerce")
    valid = pop.notna() & rate.notna() & (pop > 0)
    if not valid.any():
        return None
    return float((rate[valid] * pop[valid]).sum() / pop[valid].sum())


def _prepare_sex_table(sex_path: Path) -> pd.DataFrame:
    sex_df = pd.read_csv(sex_path)
    sex_df = _normalize_year_column(sex_df)

    sex_df[DEATHS_COL] = pd.to_numeric(sex_df.get(DEATHS_COL), errors="coerce")
    sex_df[POP_COL] = pd.to_numeric(sex_df.get(POP_COL), errors="coerce")
    sex_df[SEX_AAMR_COL] = pd.to_numeric(sex_df.get(SEX_AAMR_COL), errors="coerce")
    sex_df[SEX_COL] = sex_df[SEX_COL].astype("string").str.strip()

    # Keep only true sex strata rows; this automatically removes total/meta rows.
    sex_df = sex_df[sex_df[YEAR_COL].notna() & sex_df[SEX_COL].isin(SEX_ORDER)].copy()

    grouped = _aggregate_deaths_and_weighted_aamr(
        sex_df,
        [YEAR_COL, SEX_COL],
        deaths_col=DEATHS_COL,
        rate_col=SEX_AAMR_COL,
        pop_col=POP_COL,
    )
    return grouped


def _prepare_age_table(age_path: Path) -> pd.DataFrame:
    age_df = pd.read_csv(age_path)
    age_df = _normalize_year_column(age_df)

    age_df[DEATHS_COL] = pd.to_numeric(age_df.get(DEATHS_COL), errors="coerce")
    age_df[POP_COL] = pd.to_numeric(age_df.get(POP_COL), errors="coerce")
    age_df[AGE_AAMR_COL] = pd.to_numeric(age_df.get(AGE_AAMR_COL), errors="coerce")
    age_df[AGE_GROUP_COL] = age_df[AGE_GROUP_COL].astype("string").str.strip()

    age_df = age_df[age_df[YEAR_COL].notna() & age_df[AGE_GROUP_COL].isin(AGE_ORDER)].copy()

    grouped = _aggregate_deaths_and_weighted_aamr(
        age_df,
        [YEAR_COL, AGE_GROUP_COL],
        deaths_col=DEATHS_COL,
        rate_col=AGE_AAMR_COL,
        pop_col=POP_COL,
    )
    return grouped


def _build_outcome_yearly_table(outcome_name: str, paths: dict[str, object]) -> pd.DataFrame:
    sex_grouped = _prepare_sex_table(_resolve_sex_path(paths["sex"]))
    age_grouped = _prepare_age_table(paths["age"])

    if sex_grouped.empty:
        raise ValueError(
            f"No sex rows after filtering for {outcome_name}. "
            f"Verify 'Year', 'Sex', 'Deaths', 'Population', and '{SEX_AAMR_COL}' in the sex source."
        )

    years = sorted(
        set(sex_grouped[YEAR_COL].dropna().astype(int).tolist())
        | set(age_grouped[YEAR_COL].dropna().astype(int).tolist())
    )
    table = pd.DataFrame({"Year": years})

    for sex in SEX_ORDER:
        sex_slice = sex_grouped[sex_grouped[SEX_COL].eq(sex)].copy()
        sex_slice["Year"] = sex_slice[YEAR_COL].astype(int)
        sex_slice = sex_slice.rename(
            columns={
                "Deaths": f"{sex} Deaths",
                "AAMR": f"{sex} AAMR",
            }
        )[["Year", f"{sex} Deaths", f"{sex} AAMR"]]
        table = table.merge(sex_slice, on="Year", how="left")

    for age_group in AGE_ORDER:
        age_slice = age_grouped[age_grouped[AGE_GROUP_COL].eq(age_group)].copy()
        age_slice["Year"] = age_slice[YEAR_COL].astype(int)
        age_slice = age_slice.rename(
            columns={
                "Deaths": f"{age_group} Deaths",
                "AAMR": f"{age_group} AAMR",
            }
        )[["Year", f"{age_group} Deaths", f"{age_group} AAMR"]]
        table = table.merge(age_slice, on="Year", how="left")

    table.insert(0, "Outcome", outcome_name)
    table = table.sort_values("Year").reset_index(drop=True)

    total_row: dict[str, object] = {"Outcome": outcome_name, "Year": "Total"}
    for sex in SEX_ORDER:
        subset = sex_grouped[sex_grouped[SEX_COL].eq(sex)].copy()
        total_row[f"{sex} Deaths"] = subset["Deaths"].sum(min_count=1)
        total_row[f"{sex} AAMR"] = _weighted_rate_from_grouped(subset)

    for age_group in AGE_ORDER:
        subset = age_grouped[age_grouped[AGE_GROUP_COL].eq(age_group)].copy()
        total_row[f"{age_group} Deaths"] = subset["Deaths"].sum(min_count=1)
        total_row[f"{age_group} AAMR"] = _weighted_rate_from_grouped(subset)

    table = pd.concat([table, pd.DataFrame([total_row])], ignore_index=True)
    return table


def _fmt_two_decimals(value: float | int | None, *, with_commas: bool = False) -> str:
    if pd.isna(value):
        return ""
    if with_commas:
        return f"{float(value):,.2f}"
    return f"{float(value):.2f}"


def _format_for_export(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    for column in out.columns:
        if column not in {"Year", "Outcome"}:
            out[column] = out[column].map(lambda x: _fmt_two_decimals(x, with_commas=True))
    numeric_year = pd.to_numeric(out["Year"], errors="coerce")
    out["Year"] = numeric_year.astype("Int64").astype("string")
    out.loc[numeric_year.isna(), "Year"] = out.loc[numeric_year.isna(), "Year"].astype("string")
    return out


def _format_for_csv(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    for column in out.columns:
        if column in {"Year", "Outcome"}:
            continue
        out[column] = pd.to_numeric(out[column], errors="coerce").map(
            lambda x: _fmt_two_decimals(x, with_commas=False)
        )
    out["Year"] = out["Year"].astype("string")
    return out


def _add_publication_table(document: Document, outcome_name: str, table_df: pd.DataFrame, table_num: int) -> None:
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption_run = caption.add_run(
        f"Table {table_num}. Yearly deaths and age-adjusted mortality rates by sex and age group for {outcome_name}."
    )
    caption_run.bold = True
    caption_run.font.name = "Times New Roman"
    caption_run.font.size = Pt(10)

    display_df = _format_for_export(table_df.drop(columns=["Outcome"]))

    word_table = document.add_table(rows=2, cols=len(display_df.columns))
    word_table.style = "Table Grid"
    word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    word_table.autofit = True

    top_header = word_table.rows[0].cells
    sub_header = word_table.rows[1].cells

    _set_cell_text(top_header[0], "Year", bold=True)
    _set_cell_text(sub_header[0], "", bold=True)
    top_header[0].merge(sub_header[0])

    _set_cell_text(top_header[1], "Sex", bold=True)
    _set_cell_text(top_header[5], "Age Group", bold=True)
    top_header[1].merge(top_header[4])
    top_header[5].merge(top_header[12])

    sub_headers = [
        "Female\nDeaths",
        "Female\nAAMR",
        "Male\nDeaths",
        "Male\nAAMR",
        "<25\nDeaths",
        "<25\nAAMR",
        "25-44\nDeaths",
        "25-44\nAAMR",
        "45-64\nDeaths",
        "45-64\nAAMR",
        "65+\nDeaths",
        "65+\nAAMR",
    ]
    for idx, label in enumerate(sub_headers, start=1):
        _set_cell_text(sub_header[idx], label, bold=True, size=8)

    for row_idx in [0, 1]:
        for cell in word_table.rows[row_idx].cells:
            _set_cell_shading(cell, "D9D9D9")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for _, record in display_df.iterrows():
        row_cells = word_table.add_row().cells
        for idx, col_name in enumerate(display_df.columns):
            text = "" if pd.isna(record[col_name]) else str(record[col_name])
            align = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_text(row_cells[idx], text, bold=False, size=8, align=align)
            row_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    note = document.add_paragraph()
    note_run = note.add_run(
        "Abbreviation: AAMR, age-adjusted mortality rate per 100,000 population."
    )
    note_run.italic = True
    note_run.font.name = "Times New Roman"
    note_run.font.size = Pt(9)


def build_yearly_age_sex_tables_and_doc() -> None:
    TABLE_DIR.mkdir(parents=True, exist_ok=True)

    outcome_tables: list[pd.DataFrame] = []
    for outcome_name, paths in DATASETS.items():
        outcome_table = _build_outcome_yearly_table(outcome_name, paths)
        outcome_tables.append(outcome_table)
        outcome_slug = outcome_name.lower().replace("/", "_").replace(" ", "_")
        _format_for_csv(outcome_table).to_csv(
            TABLE_DIR / f"{outcome_slug}_yearly_age_sex_aamr_table.csv",
            index=False,
        )

    combined = pd.concat(outcome_tables, ignore_index=True)
    _format_for_csv(combined).to_csv(
        TABLE_DIR / "all_outcomes_yearly_age_sex_aamr_table.csv",
        index=False,
    )

    doc = Document()
    _configure_section(doc)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Yearly Deaths and AAMR by Sex and Age Group")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(14)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Outcomes: Pneumonia, Pneumonia/ARDS, and Pneumonia/Sepsis")
    subtitle_run.italic = True
    subtitle_run.font.name = "Times New Roman"
    subtitle_run.font.size = Pt(10)

    for idx, outcome_name in enumerate(DATASETS.keys(), start=1):
        table_df = next(df for df in outcome_tables if df["Outcome"].iloc[0] == outcome_name)
        _add_publication_table(doc, outcome_name, table_df, idx)
        if idx < len(DATASETS):
            doc.add_page_break()

    doc.save(DOC_OUTPUT_PATH)


if __name__ == "__main__":
    build_yearly_age_sex_tables_and_doc()
