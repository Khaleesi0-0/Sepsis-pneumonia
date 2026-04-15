from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
TABLE_DIR = ROOT / "results" / "tables"
CLEANED_DIR = ROOT / "data" / "cleaned"

AAMR_PATH = TABLE_DIR / "monthly_aamr_2010_2025_selected_model_table.csv"
DEATHS_VALIDATION_PATH = TABLE_DIR / "monthly_deaths_validation_2018_2019.csv"
DEATHS_PREDICTION_PATH = TABLE_DIR / "monthly_deaths_predicted_2020_2025.csv"
SELECTED_MODELS_PATH = TABLE_DIR / "monthly_aamr_validation_metrics.csv"

MONTHLY_DEATHS_FILES = {
    "Pneumonia": CLEANED_DIR / "pneumonia_month.csv",
    "Pneumonia/ARDS": CLEANED_DIR / "ards_month.csv",
    "Pneumonia/Sepsis": CLEANED_DIR / "combined_month.csv",
}

CSV_OUTPUT_PATH = TABLE_DIR / "monthly_aamr_deaths_combined_2010_2025.csv"
DOC_OUTPUT_PATH = TABLE_DIR / "monthly_aamr_deaths_combined_2010_2025.docx"

DISEASE_ORDER = ["Pneumonia", "Pneumonia/ARDS", "Pneumonia/Sepsis"]
DISEASE_NAME_MAP = {
    "ARDS (ARDS + Pneumonia)": "Pneumonia/ARDS",
    "ARDS/Pneumonia": "Pneumonia/ARDS",
    "Sepsis+ Pneumonia": "Pneumonia/Sepsis",
    "Sepsis/Pneumonia": "Pneumonia/Sepsis",
    "Combined": "Pneumonia/Sepsis",
    "combined": "Pneumonia/Sepsis",
}

OUTPUT_COLUMNS = [
    "Disease",
    "Model",
    "Validation Year",
    "Training Years",
    "Year",
    "Month",
    "Age Adjusted Rate",
    "Predicted AAMR",
    "Death",
    "Predicted Death",
    "Excess AAMR",
    "Excess Death",
]


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, *, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = 0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)


def _style_table(table) -> None:
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8 if row_idx > 0 else 9)
        if row_idx == 0:
            for cell in row.cells:
                _set_cell_shading(cell, "D9E2F3")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def _set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)


def _month_number_from_code(month_code: pd.Series) -> pd.Series:
    return pd.to_numeric(
        month_code.astype("string").str.extract(r"/(\d{2})$", expand=False),
        errors="coerce",
    )


def _month_number_from_label(month_label: pd.Series) -> pd.Series:
    normalized = (
        month_label.astype("string")
        .str.replace(".", "", regex=False)
        .str.strip()
    )
    parsed = pd.to_datetime(normalized, format="%b, %Y", errors="coerce")
    return parsed.dt.month


def _normalize_disease(series: pd.Series) -> pd.Series:
    return series.astype("string").str.strip().replace(DISEASE_NAME_MAP)


def _format_num(series: pd.Series) -> pd.Series:
    numeric = pd.to_numeric(series, errors="coerce")
    return numeric.map(lambda x: "" if pd.isna(x) else f"{x:.2f}")


def _build_combined_df() -> pd.DataFrame:
    if not AAMR_PATH.exists():
        raise FileNotFoundError(f"Missing source table: {AAMR_PATH}")
    if not DEATHS_VALIDATION_PATH.exists():
        raise FileNotFoundError(f"Missing source table: {DEATHS_VALIDATION_PATH}")
    if not DEATHS_PREDICTION_PATH.exists():
        raise FileNotFoundError(f"Missing source table: {DEATHS_PREDICTION_PATH}")
    if not SELECTED_MODELS_PATH.exists():
        raise FileNotFoundError(f"Missing selected-model table: {SELECTED_MODELS_PATH}")
    for disease, path in MONTHLY_DEATHS_FILES.items():
        if not path.exists():
            raise FileNotFoundError(f"Missing monthly deaths source for {disease}: {path}")

    selected_models = pd.read_csv(SELECTED_MODELS_PATH)
    selected_models["Disease"] = _normalize_disease(selected_models["Disease"])
    selected_models["Selected Model"] = selected_models["Selected Model"].astype("string").str.strip()
    selected_models = selected_models[["Disease", "Selected Model"]].drop_duplicates()

    aamr = pd.read_csv(AAMR_PATH)
    aamr["Disease"] = _normalize_disease(aamr["Disease"])
    aamr["Year"] = pd.to_numeric(aamr["Year"], errors="coerce")
    aamr["__month_num"] = _month_number_from_label(aamr["Month"])
    aamr["Month Code"] = (
        aamr["Year"].astype("Int64").astype("string").str.replace("<NA>", "", regex=False)
        + "/"
        + aamr["__month_num"].astype("Int64").astype("string").str.zfill(2)
    )
    aamr["Excess AAMR"] = aamr.get("Actual - Predicted", pd.NA)
    aamr = aamr.merge(selected_models, on="Disease", how="left")
    aamr = aamr[aamr["Model"].astype("string").str.strip().eq(aamr["Selected Model"].astype("string").str.strip())].copy()
    aamr = aamr.drop(columns=["Selected Model"])

    deaths_val = pd.read_csv(DEATHS_VALIDATION_PATH, dtype={"Year": "string", "Month Code": "string"})
    deaths_pred = pd.read_csv(DEATHS_PREDICTION_PATH, dtype={"Year": "string", "Month Code": "string"})
    deaths = pd.concat([deaths_val, deaths_pred], ignore_index=True, sort=False)
    deaths["Disease"] = _normalize_disease(deaths["Disease"])
    deaths["Year"] = pd.to_numeric(deaths["Year"], errors="coerce")
    deaths["__month_num"] = _month_number_from_code(deaths["Month Code"])
    deaths = deaths.rename(
        columns={
            "Deaths": "Death",
            "Predicted Deaths": "Predicted Death",
            "Actual - Predicted Deaths": "Excess Death",
        }
    )
    deaths = deaths.merge(selected_models, on="Disease", how="left")
    deaths = deaths[
        deaths["Model"].astype("string").str.strip().eq(
            deaths["Selected Model"].astype("string").str.strip()
        )
    ].copy()
    deaths = deaths.drop(columns=["Selected Model"])

    actual_death_frames = []
    for disease, path in MONTHLY_DEATHS_FILES.items():
        src = pd.read_csv(path, dtype={"Year": "string", "Month Code": "string"})
        src["Disease"] = disease
        src["Year"] = pd.to_numeric(src["Year"], errors="coerce")
        src["Month Code"] = src["Month Code"].astype("string").str.strip()
        src["Actual Death"] = pd.to_numeric(src["Deaths"], errors="coerce")
        src = src.dropna(subset=["Year", "Month Code", "Actual Death"]).copy()
        src["Year"] = src["Year"].astype(int)
        actual_death_frames.append(src[["Disease", "Year", "Month Code", "Actual Death"]])
    actual_deaths = pd.concat(actual_death_frames, ignore_index=True)

    merge_keys = ["Disease", "Year", "Month Code"]
    merged = aamr.merge(
        deaths[
            merge_keys
            + [
                "Death",
                "Predicted Death",
                "Excess Death",
                "__month_num",
            ]
        ],
        on=merge_keys,
        how="left",
        suffixes=("", "_death"),
    )
    merged = merged.merge(actual_deaths, on=merge_keys, how="left")
    merged["Death"] = pd.to_numeric(merged["Death"], errors="coerce").combine_first(
        pd.to_numeric(merged["Actual Death"], errors="coerce")
    )
    merged = merged.drop(columns=["Actual Death"])
    merged["__month_num"] = merged["__month_num"].fillna(merged.get("__month_num_death"))
    merged = merged.drop(columns=[c for c in ["__month_num_death"] if c in merged.columns])

    merged["Disease"] = pd.Categorical(merged["Disease"], categories=DISEASE_ORDER, ordered=True)
    merged = merged.sort_values(["Disease", "Year", "__month_num"], na_position="last").reset_index(drop=True)

    for col in [
        "Age Adjusted Rate",
        "Predicted AAMR",
        "Excess AAMR",
        "Death",
        "Predicted Death",
        "Excess Death",
    ]:
        if col in merged.columns:
            merged[col] = _format_num(merged[col])

    merged["Validation Year"] = merged["Validation Year"].astype("string").replace("<NA>", "")
    merged["Training Years"] = merged["Training Years"].astype("string").replace("<NA>", "")
    merged["Year"] = merged["Year"].astype("Int64").astype("string")

    out = merged[OUTPUT_COLUMNS].copy()
    out["Disease"] = out["Disease"].astype("string")
    return out


def _add_df_table(document: Document, title: str, df: pd.DataFrame) -> None:
    heading = document.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    table = document.add_table(rows=1, cols=len(df.columns))
    table.autofit = True

    for i, col in enumerate(df.columns):
        _set_cell_text(table.rows[0].cells[i], col, bold=True, size=9)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            _set_cell_text(cells[i], row[col], size=8)

    _style_table(table)
    document.add_paragraph("")


def _export_doc(df: pd.DataFrame) -> None:
    doc = Document()
    _set_landscape(doc.sections[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = 1
    run = title.add_run("Combined Monthly AAMR and Death Table (2010-2025)")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    subtitle = doc.add_paragraph()
    subtitle.alignment = 1
    run = subtitle.add_run("Observed, predicted, and excess values by disease and month")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    for idx, disease in enumerate(DISEASE_ORDER):
        disease_df = df[df["Disease"].eq(disease)].copy()
        _add_df_table(doc, disease, disease_df)
        if idx < len(DISEASE_ORDER) - 1:
            doc.add_page_break()

    DOC_OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOC_OUTPUT_PATH)


def build_and_export() -> None:
    TABLE_DIR.mkdir(parents=True, exist_ok=True)
    out = _build_combined_df()
    out.to_csv(CSV_OUTPUT_PATH, index=False)
    _export_doc(out)


if __name__ == "__main__":
    build_and_export()
