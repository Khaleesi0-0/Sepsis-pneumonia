from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
TABLE_DIR = ROOT / "results" / "tables"
DOC_PATH = TABLE_DIR / "monthly_deaths_validation_and_predicted_tables.docx"

TABLE_SOURCES = [
    TABLE_DIR / "monthly_deaths_validation_2018_2019.csv",
    TABLE_DIR / "monthly_deaths_predicted_2020_2025.csv",
]

DISEASE_ORDER = ["Pneumonia", "Pneumonia/ARDS", "Pneumonia/Sepsis"]
DISEASE_NAME_MAP = {
    "Combined": "Pneumonia/Sepsis",
    "combined": "Pneumonia/Sepsis",
    "Sepsis+ Pneumonia": "Pneumonia/Sepsis",
    "Sepsis/Pneumonia": "Pneumonia/Sepsis",
    "ARDS (ARDS + Pneumonia)": "Pneumonia/ARDS",
    "ARDS/Pneumonia": "Pneumonia/ARDS",
}


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, *, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = 0
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)


def _style_table(table) -> None:
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8 if row_idx > 0 else 9)
        if row_idx == 0:
            for cell in row.cells:
                _set_cell_shading(cell, "D9E2F3")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def _set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)


def _add_dataframe_table(document: Document, title: str, df: pd.DataFrame) -> None:
    heading = document.add_paragraph()
    heading_run = heading.add_run(title)
    heading_run.bold = True
    heading_run.font.name = "Times New Roman"
    heading_run.font.size = Pt(11)

    table = document.add_table(rows=1, cols=len(df.columns))
    table.autofit = True

    for col_idx, column_name in enumerate(df.columns):
        _set_cell_text(table.rows[0].cells[col_idx], column_name, bold=True, size=9)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for col_idx, value in enumerate(row.tolist()):
            _set_cell_text(cells[col_idx], "" if pd.isna(value) else value, size=8)

    _style_table(table)
    document.add_paragraph("")


def _month_number_from_code(month_code: pd.Series) -> pd.Series:
    return pd.to_numeric(
        month_code.astype("string").str.extract(r"/(\d{2})$", expand=False),
        errors="coerce",
    )


def _normalize_and_sort(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    if "Disease" in out.columns:
        out["Disease"] = out["Disease"].astype("string").str.strip().replace(DISEASE_NAME_MAP)
        out["Disease"] = pd.Categorical(out["Disease"], categories=DISEASE_ORDER, ordered=True)

    if "Year" in out.columns:
        out["__year_num"] = pd.to_numeric(out["Year"], errors="coerce")
    else:
        out["__year_num"] = pd.NA

    if "Month Code" in out.columns:
        out["__month_num"] = _month_number_from_code(out["Month Code"])
    else:
        out["__month_num"] = pd.NA

    sort_cols = []
    if "Disease" in out.columns:
        sort_cols.append("Disease")
    sort_cols.extend(["__year_num", "__month_num"])
    out = out.sort_values(sort_cols, na_position="last").reset_index(drop=True)
    return out.drop(columns=["__year_num", "__month_num"], errors="ignore")


def export_tables() -> None:
    document = Document()
    _set_landscape(document.sections[0])

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(9)

    title = document.add_paragraph()
    title.alignment = 1
    run = title.add_run("Monthly Deaths Validation and Predicted Tables")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    subtitle = document.add_paragraph()
    subtitle.alignment = 1
    run = subtitle.add_run("Publication-style export")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    for idx, table_path in enumerate(TABLE_SOURCES):
        if not table_path.exists():
            continue
        df = _normalize_and_sort(pd.read_csv(table_path))
        title = table_path.stem.replace("_", " ").title()

        if "Disease" not in df.columns:
            _add_dataframe_table(document, title, df)
        else:
            section_heading = document.add_paragraph()
            run = section_heading.add_run(title)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

            for disease in DISEASE_ORDER:
                disease_df = df[df["Disease"].astype("string").eq(disease)].copy()
                if disease_df.empty:
                    continue
                _add_dataframe_table(document, disease, disease_df)

        if idx < len(TABLE_SOURCES) - 1:
            document.add_page_break()

    DOC_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(DOC_PATH)


if __name__ == "__main__":
    export_tables()
