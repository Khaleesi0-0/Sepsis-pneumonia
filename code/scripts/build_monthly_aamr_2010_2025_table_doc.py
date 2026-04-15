from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
CLEANED_DIR = ROOT / "data" / "cleaned"
TABLE_DIR = ROOT / "results" / "tables"

VALIDATION_PATH = TABLE_DIR / "monthly_aamr_validation_2018_2019.csv"
PREDICTION_PATH = TABLE_DIR / "monthly_aamr_predicted_2020_2025.csv"
METRICS_PATH = TABLE_DIR / "monthly_aamr_validation_metrics.csv"

CSV_OUTPUT_PATH = TABLE_DIR / "monthly_aamr_2010_2025_selected_model_table.csv"
DOC_OUTPUT_PATH = TABLE_DIR / "monthly_aamr_2010_2025_selected_model_table.docx"

YEAR_COL = "Year"
MONTH_COL = "Month"
MONTH_CODE_COL = "Month Code"
RATE_COL = "Age Adjusted Rate"
PREDICTED_COL = "Predicted AAMR"
DIFF_COL = "Actual - Predicted"

YEAR_START = 2010
YEAR_END = 2025

DATASETS = {
    "Pneumonia": CLEANED_DIR / "pneumonia_month.csv",
    "Pneumonia/ARDS": CLEANED_DIR / "ards_month.csv",
    "Pneumonia/Sepsis": CLEANED_DIR / "combined_month.csv",
}

DISEASE_NAME_MAP = {
    "ARDS (ARDS + Pneumonia)": "Pneumonia/ARDS",
    "Sepsis+ Pneumonia": "Pneumonia/Sepsis",
    "ARDS/Pneumonia": "Pneumonia/ARDS",
    "Sepsis/Pneumonia": "Pneumonia/Sepsis",
}

OUTPUT_COLUMNS = [
    "Disease",
    "Model",
    "Validation Year",
    "Training Years",
    "Year",
    "Month",
    "Age Adjusted Rate",
    "Predicted AAMR",
    "Actual - Predicted",
]


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, *, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = 0
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)


def _style_table(table) -> None:
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8 if row_idx > 0 else 9)
        if row_idx == 0:
            for cell in row.cells:
                _set_cell_shading(cell, "D9E2F3")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def _set_landscape(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)


def _month_number_from_code(month_code: pd.Series) -> pd.Series:
    return pd.to_numeric(
        month_code.astype("string").str.extract(r"/(\d{2})$", expand=False),
        errors="coerce",
    )


def _normalize_year(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out[YEAR_COL] = (
        out[YEAR_COL]
        .astype("string")
        .str.strip()
        .str.replace(r"\s*\(provisional\)$", "", regex=True)
    )
    out[YEAR_COL] = pd.to_numeric(out[YEAR_COL], errors="coerce")
    return out


def _read_actual_monthly() -> pd.DataFrame:
    frames: list[pd.DataFrame] = []
    for disease, path in DATASETS.items():
        df = pd.read_csv(path, dtype={YEAR_COL: "string", MONTH_CODE_COL: "string"})
        df = _normalize_year(df)
        df[RATE_COL] = pd.to_numeric(df[RATE_COL], errors="coerce")
        df["__month_num"] = _month_number_from_code(df[MONTH_CODE_COL])
        df = df[
            df[YEAR_COL].between(YEAR_START, YEAR_END)
            & df["__month_num"].between(1, 12)
            & df[RATE_COL].notna()
        ].copy()
        df["Disease"] = disease
        frames.append(df[["Disease", YEAR_COL, MONTH_COL, MONTH_CODE_COL, "__month_num", RATE_COL]])
    out = pd.concat(frames, ignore_index=True)
    out[YEAR_COL] = out[YEAR_COL].astype(int)
    out["__month_num"] = out["__month_num"].astype(int)
    return out


def _read_selected_model_lookup() -> pd.DataFrame:
    metrics = pd.read_csv(METRICS_PATH)
    metrics["Disease"] = metrics["Disease"].astype("string").str.strip().replace(DISEASE_NAME_MAP)
    metrics["Final Selected Model"] = metrics["Selected Model"].astype("string").str.strip()
    return metrics[["Disease", "Final Selected Model"]].drop_duplicates()


def _read_prediction_rows() -> pd.DataFrame:
    validation = pd.read_csv(
        VALIDATION_PATH,
        dtype={YEAR_COL: "string", MONTH_CODE_COL: "string"},
    )
    validation["Disease"] = validation["Disease"].astype("string").str.strip().replace(DISEASE_NAME_MAP)
    validation[YEAR_COL] = pd.to_numeric(validation[YEAR_COL], errors="coerce")
    validation["Validation Year"] = pd.to_numeric(validation["Validation Year"], errors="coerce").astype("Int64").astype("string")
    validation["Training Years"] = validation["Training Years"].astype("string")
    validation[PREDICTED_COL] = pd.to_numeric(validation[PREDICTED_COL], errors="coerce")
    validation[DIFF_COL] = pd.to_numeric(validation[DIFF_COL], errors="coerce")
    validation = validation.rename(columns={"Selected Model": "Model"})

    prediction = pd.read_csv(
        PREDICTION_PATH,
        dtype={YEAR_COL: "string", MONTH_CODE_COL: "string"},
    )
    prediction["Disease"] = prediction["Disease"].astype("string").str.strip().replace(DISEASE_NAME_MAP)
    prediction[YEAR_COL] = pd.to_numeric(prediction[YEAR_COL], errors="coerce")
    prediction["Validation Year"] = pd.NA
    prediction["Training Years"] = "2010-2019"
    prediction[PREDICTED_COL] = pd.to_numeric(prediction[PREDICTED_COL], errors="coerce")
    prediction[DIFF_COL] = pd.to_numeric(prediction[DIFF_COL], errors="coerce")
    prediction = prediction.rename(columns={"Selected Model": "Model"})

    combined = pd.concat([validation, prediction], ignore_index=True, sort=False)
    combined["__month_num"] = _month_number_from_code(combined[MONTH_CODE_COL])
    combined = combined.dropna(subset=["Disease", YEAR_COL, "__month_num", "Model"])
    combined[YEAR_COL] = combined[YEAR_COL].astype(int)
    combined["__month_num"] = combined["__month_num"].astype(int)
    combined = combined[
        [
            "Disease",
            "Model",
            "Validation Year",
            "Training Years",
            YEAR_COL,
            MONTH_CODE_COL,
            "__month_num",
            PREDICTED_COL,
            DIFF_COL,
        ]
    ].copy()
    return combined


def _format_numeric(series: pd.Series) -> pd.Series:
    numeric = pd.to_numeric(series, errors="coerce")
    return numeric.map(lambda x: "" if pd.isna(x) else f"{x:.1f}")


def build_monthly_aamr_2010_2025_table() -> pd.DataFrame:
    actual = _read_actual_monthly()
    predicted = _read_prediction_rows()
    selected_model = _read_selected_model_lookup()

    # Keep only predictions from the final selected model for each disease.
    predicted = predicted.merge(selected_model, on="Disease", how="left")
    predicted = predicted[
        predicted["Model"].astype("string").str.strip().eq(
            predicted["Final Selected Model"].astype("string").str.strip()
        )
    ].copy()
    predicted = predicted.drop(columns=["Final Selected Model"])

    merged = actual.merge(
        predicted,
        on=["Disease", YEAR_COL, MONTH_CODE_COL, "__month_num"],
        how="left",
    )
    merged = merged.merge(selected_model, on="Disease", how="left")
    merged["Model"] = merged["Final Selected Model"]
    merged = merged.drop(columns=["Final Selected Model"])

    merged["Validation Year"] = merged["Validation Year"].astype("string").replace("<NA>", "")
    merged["Training Years"] = merged["Training Years"].astype("string").replace("<NA>", "")

    out = merged.rename(columns={RATE_COL: "Age Adjusted Rate"})[
        [
            "Disease",
            "Model",
            "Validation Year",
            "Training Years",
            YEAR_COL,
            MONTH_COL,
            "Age Adjusted Rate",
            PREDICTED_COL,
            DIFF_COL,
            "__month_num",
        ]
    ].copy()

    out["Age Adjusted Rate"] = _format_numeric(out["Age Adjusted Rate"])
    out[PREDICTED_COL] = _format_numeric(out[PREDICTED_COL])
    out[DIFF_COL] = _format_numeric(out[DIFF_COL])

    out = out.sort_values(["Disease", YEAR_COL, "__month_num"]).reset_index(drop=True)
    out = out.drop(columns="__month_num")
    out = out.rename(
        columns={
            YEAR_COL: "Year",
            MONTH_COL: "Month",
        }
    )
    return out[OUTPUT_COLUMNS]


def _add_disease_table(document: Document, disease_df: pd.DataFrame, disease_name: str) -> None:
    heading = document.add_paragraph()
    run = heading.add_run(disease_name)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    table = document.add_table(rows=1, cols=len(OUTPUT_COLUMNS))
    for idx, col in enumerate(OUTPUT_COLUMNS):
        _set_cell_text(table.rows[0].cells[idx], col, bold=True, size=9)

    for _, row in disease_df.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(OUTPUT_COLUMNS):
            _set_cell_text(cells[idx], row[col], size=8)

    _style_table(table)
    document.add_paragraph("")


def export_monthly_table_doc(table_df: pd.DataFrame) -> None:
    doc = Document()
    _set_landscape(doc)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = 1
    run = title.add_run("Monthly AAMR Table (2010-2025)")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    subtitle = doc.add_paragraph()
    subtitle.alignment = 1
    run = subtitle.add_run("Observed and selected-model predicted monthly age-adjusted mortality rates")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    for disease in DATASETS.keys():
        disease_df = table_df[table_df["Disease"].eq(disease)].copy()
        _add_disease_table(doc, disease_df, disease)
        if disease != list(DATASETS.keys())[-1]:
            doc.add_page_break()

    DOC_OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOC_OUTPUT_PATH)


def build_and_export() -> None:
    TABLE_DIR.mkdir(parents=True, exist_ok=True)
    out = build_monthly_aamr_2010_2025_table()
    out.to_csv(CSV_OUTPUT_PATH, index=False)
    export_monthly_table_doc(out)


if __name__ == "__main__":
    build_and_export()
